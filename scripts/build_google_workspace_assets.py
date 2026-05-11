from __future__ import annotations

import html
import re
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "exports" / "google-workspace"
CEO = ROOT / "docs" / "ceo-framework"
COM = ROOT / "docs" / "comercial"


DOCS = [
    (
        "01_Humanio_CEO_Framework_Manual_Maestro.docx",
        "Humanio CEO Framework - Manual Maestro",
        [
            CEO / "00_MANIFIESTO_CEO.md",
            CEO / "01_FRAMEWORK_CEO.md",
            CEO / "10_FLUJO_VISUAL_CEO.md",
        ],
    ),
    (
        "02_Humanio_CEO_Playbook_Implementacion.docx",
        "Humanio CEO Framework - Playbook de Implementacion",
        [
            CEO / "06_PLAN_IMPLEMENTACION.md",
            CEO / "07_EVALUACION_QA.md",
            CEO / "08_GOBIERNO_MEJORA.md",
        ],
    ),
    (
        "03_Humanio_CEO_Plantilla_Diagnostico.docx",
        "Humanio CEO Framework - Plantilla de Diagnostico",
        [CEO / "02_DIAGNOSTICO_CLIENTE.md", COM / "02_GUIA_SESION_DIAGNOSTICO.md"],
    ),
    (
        "04_Humanio_CEO_Plantilla_Blueprint_Agente.docx",
        "Humanio CEO Framework - Plantilla Blueprint del Agente",
        [CEO / "03_BLUEPRINT_AGENTE.md"],
    ),
    (
        "05_Humanio_CEO_Ecosistema_Orquestacion.docx",
        "Humanio CEO Framework - Ecosistema y Orquestacion",
        [CEO / "04_MAPA_ECOSISTEMA.md", CEO / "05_ORQUESTACION_AGENTES.md"],
    ),
    (
        "06_Humanio_CEO_Entregable_Ejecutivo_Cliente.docx",
        "Humanio CEO Framework - Entregable Ejecutivo para Cliente",
        [CEO / "09_ENTREGABLE_EJECUTIVO_CLIENTE.md"],
    ),
    (
        "07_Humanio_CEO_Narrativa_Comercial.docx",
        "Humanio CEO Framework - Narrativa Comercial",
        [COM / "00_NARRATIVA_COMERCIAL.md"],
    ),
]


def style_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    for name, size, color in [
        ("Title", 22, RGBColor(26, 45, 79)),
        ("Heading 1", 16, RGBColor(26, 45, 79)),
        ("Heading 2", 13, RGBColor(44, 83, 100)),
        ("Heading 3", 11.5, RGBColor(60, 60, 60)),
    ]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = color


def add_markdown(document: Document, markdown: str) -> None:
    in_code = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if len(table_rows) < 2:
            table_rows = []
            return
        rows = [r for r in table_rows if not all(re.fullmatch(r"-+", c.strip()) for c in r)]
        if rows:
            table = document.add_table(rows=len(rows), cols=max(len(r) for r in rows))
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    if j < len(table.rows[i].cells):
                        table.rows[i].cells[j].text = cell.strip()
                        for p in table.rows[i].cells[j].paragraphs:
                            for run in p.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(9)
                        if i == 0:
                            for p in table.rows[i].cells[j].paragraphs:
                                for run in p.runs:
                                    run.bold = True
            document.add_paragraph()
        table_rows = []

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.strip().startswith("```"):
            in_code = not in_code
            flush_table()
            continue
        if in_code:
            continue
        if "|" in line and line.strip().startswith("|"):
            table_rows.append([c.strip() for c in line.strip().strip("|").split("|")])
            continue
        flush_table()
        if not line.strip():
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
        elif line.startswith("- [ ] "):
            document.add_paragraph("☐ " + line[6:].strip())
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            p = document.add_paragraph()
            text = re.sub(r"\*\*(.*?)\*\*", r"\1", line)
            p.add_run(text)
    flush_table()


def build_docx() -> None:
    for filename, title, sources in DOCS:
        document = Document()
        style_doc(document)
        title_p = document.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(title)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(26, 45, 79)
        subtitle = document.add_paragraph("CEO: Contexto, Ecosistema y Orquestacion")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.name = "Arial"
        subtitle.runs[0].font.size = Pt(11)
        document.add_paragraph()
        for source in sources:
            add_markdown(document, source.read_text(encoding="utf-8"))
            if source != sources[-1]:
                document.add_page_break()
        document.save(OUT / filename)


def add_sheet(wb: Workbook, title: str, headers: list[str], rows: list[list[str]]) -> None:
    ws = wb.create_sheet(title)
    ws.append(headers)
    for row in rows:
        ws.append(row)
    header_fill = PatternFill("solid", fgColor="1A2D4F")
    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for col in range(1, ws.max_column + 1):
        ws.column_dimensions[get_column_letter(col)].width = 24
    ws.freeze_panes = "A2"


def build_xlsx() -> None:
    wb = Workbook()
    wb.remove(wb.active)
    add_sheet(
        wb,
        "00 Indice",
        ["Documento", "Uso", "Responsable", "Estado"],
        [
            ["Diagnostico", "Levantar oportunidad y viabilidad", "Consultor", "Pendiente"],
            ["Priorizacion", "Elegir primer agente", "Consultor + Cliente", "Pendiente"],
            ["Blueprint", "Definir agente", "Consultor", "Pendiente"],
            ["Ecosistema", "Mapear canales, sistemas y datos", "Consultor + TI", "Pendiente"],
            ["Orquestacion", "Definir estados, handoffs y reglas", "Consultor", "Pendiente"],
            ["QA", "Probar calidad y riesgos", "Consultor + Operacion", "Pendiente"],
            ["Produccion", "Aprobar salida controlada", "Responsable proyecto", "Pendiente"],
            ["Gobierno", "Medir y mejorar", "Direccion + Operacion", "Pendiente"],
        ],
    )
    add_sheet(
        wb,
        "01 Diagnostico",
        ["Bloque", "Pregunta", "Respuesta", "Evidencia", "Riesgo"],
        [
            ["Negocio", "Que problema se quiere resolver?", "", "", ""],
            ["Resultado", "Que indicador debe mejorar?", "", "", ""],
            ["Usuario", "Quien usara o recibira al agente?", "", "", ""],
            ["Proceso", "Cuales son los pasos actuales?", "", "", ""],
            ["Datos", "Que fuentes existen?", "", "", ""],
            ["Restricciones", "Que no puede hacer el agente?", "", "", ""],
        ],
    )
    add_sheet(
        wb,
        "02 Priorizacion",
        ["Agente candidato", "Valor", "Viabilidad", "Riesgo invertido", "Datos", "Urgencia", "Total", "Decision"],
        [
            ["", "1-5", "1-5", "1-5", "1-5", "1-5", "=SUM(B2:F2)", ""],
            ["", "1-5", "1-5", "1-5", "1-5", "1-5", "=SUM(B3:F3)", ""],
            ["", "1-5", "1-5", "1-5", "1-5", "1-5", "=SUM(B4:F4)", ""],
        ],
    )
    add_sheet(
        wb,
        "03 Blueprint",
        ["Campo", "Definicion"],
        [
            ["Nombre del agente", ""],
            ["Rol", ""],
            ["Proposito", ""],
            ["Usuario principal", ""],
            ["Canal principal", ""],
            ["Resultado esperado", ""],
            ["Indicador principal", ""],
            ["Capacidades", ""],
            ["Limites", ""],
            ["Reglas de escalamiento", ""],
        ],
    )
    add_sheet(
        wb,
        "04 Ecosistema",
        ["Elemento", "Tipo", "Uso", "Duenio", "Estado", "Riesgo"],
        [
            ["WhatsApp", "Canal", "Atencion y seguimiento", "", "Pendiente", ""],
            ["CRM", "Sistema", "Contactos y pipeline", "", "Pendiente", ""],
            ["Calendario", "Sistema", "Disponibilidad y citas", "", "Pendiente", ""],
            ["Base de conocimiento", "Datos", "Respuestas y politicas", "", "Pendiente", ""],
        ],
    )
    add_sheet(
        wb,
        "05 Orquestacion",
        ["Estado", "Entrada", "Decision", "Accion", "Escalamiento", "Log requerido"],
        [
            ["Nuevo", "Mensaje/evento", "Clasificar intencion", "Responder o pedir datos", "No", "Fecha, usuario, intencion"],
            ["Requiere humano", "Caso sensible", "Escalar", "Enviar resumen", "Si", "Motivo, resumen, urgencia"],
            ["Error", "Falla herramienta", "Modo degradado", "Informar y registrar", "Segun impacto", "Error, herramienta, accion"],
        ],
    )
    add_sheet(
        wb,
        "06 QA",
        ["Caso", "Entrada", "Resultado esperado", "Resultado real", "Calificacion 1-5", "Accion"],
        [
            ["Caso ideal", "", "", "", "", ""],
            ["Caso incompleto", "", "", "", "", ""],
            ["Caso ambiguo", "", "", "", "", ""],
            ["Caso sensible", "", "", "", "", ""],
            ["Herramienta caida", "", "", "", "", ""],
        ],
    )
    add_sheet(
        wb,
        "07 Produccion",
        ["Categoria", "Checklist", "Estado", "Responsable", "Notas"],
        [
            ["Contexto", "Objetivo de negocio definido", "Pendiente", "", ""],
            ["Ecosistema", "Integraciones minimas probadas", "Pendiente", "", ""],
            ["Orquestacion", "Reglas de escalamiento probadas", "Pendiente", "", ""],
            ["QA", "No hay errores criticos abiertos", "Pendiente", "", ""],
            ["Produccion", "Plan de reversa definido", "Pendiente", "", ""],
        ],
    )
    add_sheet(
        wb,
        "08 Gobierno",
        ["Metrica", "Frecuencia", "Responsable", "Valor actual", "Meta", "Accion"],
        [
            ["Conversaciones atendidas", "Semanal", "", "", "", ""],
            ["Resolucion automatica", "Semanal", "", "", "", ""],
            ["Escalamiento correcto", "Semanal", "", "", "", ""],
            ["Errores criticos", "Diaria al inicio", "", "", "0", ""],
            ["Satisfaccion", "Mensual", "", "", "", ""],
        ],
    )
    wb.save(OUT / "Humanio_CEO_Workbook_Operativo.xlsx")


SLIDES = [
    ("Humanio CEO Framework", ["Contexto, Ecosistema y Orquestacion para agentes de IA confiables."]),
    ("El problema", ["Muchas iniciativas empiezan por la herramienta.", "Falta contexto, integracion, escalamiento y medicion."]),
    ("La tesis", ["Un agente serio no solo responde.", "Opera dentro de un sistema de trabajo real."]),
    ("Modelo CEO", ["Contexto: que debe entender.", "Ecosistema: donde vive y con que se conecta.", "Orquestacion: como se coordina y mide."]),
    ("Contexto", ["Objetivo de negocio.", "Usuario final.", "Conocimiento, reglas, limites y criterios de exito."]),
    ("Ecosistema", ["Canales, CRM, calendario, datos y personas.", "Riesgos, permisos e integraciones."]),
    ("Orquestacion", ["Estados del proceso.", "Handoffs.", "Logs, metricas y modo degradado."]),
    ("Proceso de trabajo", ["Diagnostico.", "Diseno CEO.", "Blueprint.", "Prototipo.", "QA.", "Produccion.", "Mejora continua."]),
    ("Entregables", ["Diagnostico, blueprint y mapa de ecosistema.", "Flujo de orquestacion, QA y gobierno."]),
    ("Resultado esperado", ["Un agente con objetivo, limites, integraciones, medicion y supervision."]),
    ("Primer paso", ["Sesion de diagnostico CEO para elegir el primer caso viable."]),
]


def tx_body(lines: list[str], x: int, y: int, cx: int, cy: int, font_size: int = 2400) -> str:
    paragraphs = []
    for line in lines:
        paragraphs.append(
            f"""
            <a:p>
              <a:pPr marL=\"342900\" indent=\"-171450\"><a:buChar char=\"•\"/></a:pPr>
              <a:r><a:rPr lang=\"es-MX\" sz=\"{font_size}\"><a:solidFill><a:srgbClr val=\"263238\"/></a:solidFill></a:rPr><a:t>{escape(line)}</a:t></a:r>
              <a:endParaRPr lang=\"es-MX\" sz=\"{font_size}\"/>
            </a:p>"""
        )
    text = "\n".join(paragraphs)
    return f"""
    <p:sp>
      <p:nvSpPr><p:cNvPr id=\"3\" name=\"Body\"/><p:cNvSpPr txBox=\"1\"/><p:nvPr/></p:nvSpPr>
      <p:spPr><a:xfrm><a:off x=\"{x}\" y=\"{y}\"/><a:ext cx=\"{cx}\" cy=\"{cy}\"/></a:xfrm><a:prstGeom prst=\"rect\"><a:avLst/></a:prstGeom><a:noFill/></p:spPr>
      <p:txBody><a:bodyPr wrap=\"square\"/><a:lstStyle/>{text}</p:txBody>
    </p:sp>"""


def title_shape(title: str) -> str:
    return f"""
    <p:sp>
      <p:nvSpPr><p:cNvPr id=\"2\" name=\"Title\"/><p:cNvSpPr txBox=\"1\"/><p:nvPr/></p:nvSpPr>
      <p:spPr><a:xfrm><a:off x=\"685800\" y=\"457200\"/><a:ext cx=\"7772400\" cy=\"914400\"/></a:xfrm><a:prstGeom prst=\"rect\"><a:avLst/></a:prstGeom><a:noFill/></p:spPr>
      <p:txBody><a:bodyPr wrap=\"square\"/><a:lstStyle/>
        <a:p><a:r><a:rPr lang=\"es-MX\" sz=\"3800\" b=\"1\"><a:solidFill><a:srgbClr val=\"1A2D4F\"/></a:solidFill></a:rPr><a:t>{escape(title)}</a:t></a:r><a:endParaRPr lang=\"es-MX\" sz=\"3800\"/></a:p>
      </p:txBody>
    </p:sp>"""


def slide_xml(title: str, bullets: list[str], idx: int) -> str:
    accent = "2C5364" if idx % 2 else "1A2D4F"
    return f"""<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>
<p:sld xmlns:a=\"http://schemas.openxmlformats.org/drawingml/2006/main\" xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\" xmlns:p=\"http://schemas.openxmlformats.org/presentationml/2006/main\">
  <p:cSld><p:bg><p:bgPr><a:solidFill><a:srgbClr val=\"F7F9FB\"/></a:solidFill><a:effectLst/></p:bgPr></p:bg><p:spTree>
    <p:nvGrpSpPr><p:cNvPr id=\"1\" name=\"\"/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr>
    <p:grpSpPr><a:xfrm><a:off x=\"0\" y=\"0\"/><a:ext cx=\"0\" cy=\"0\"/><a:chOff x=\"0\" y=\"0\"/><a:chExt cx=\"0\" cy=\"0\"/></a:xfrm></p:grpSpPr>
    <p:sp>
      <p:nvSpPr><p:cNvPr id=\"10\" name=\"Accent\"/><p:cNvSpPr/><p:nvPr/></p:nvSpPr>
      <p:spPr><a:xfrm><a:off x=\"0\" y=\"0\"/><a:ext cx=\"274320\" cy=\"6858000\"/></a:xfrm><a:prstGeom prst=\"rect\"><a:avLst/></a:prstGeom><a:solidFill><a:srgbClr val=\"{accent}\"/></a:solidFill><a:ln><a:noFill/></a:ln></p:spPr>
    </p:sp>
    {title_shape(title)}
    {tx_body(bullets, 914400, 1600200, 7315200, 3429000)}
  </p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr>
</p:sld>"""


def build_pptx() -> None:
    path = OUT / "Humanio_CEO_Presentacion_Comercial.pptx"
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as z:
        overrides = "\n".join(
            [f'<Override PartName="/ppt/slides/slide{i}.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.slide+xml"/>' for i in range(1, len(SLIDES) + 1)]
        )
        z.writestr(
            "[Content_Types].xml",
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/>
{overrides}
</Types>""",
        )
        z.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="ppt/presentation.xml"/>
</Relationships>""",
        )
        slide_ids = "\n".join([f'<p:sldId id="{255+i}" r:id="rId{i}"/>' for i in range(1, len(SLIDES) + 1)])
        z.writestr(
            "ppt/presentation.xml",
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<p:presentation xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main">
<p:sldSz cx="9144000" cy="5143500" type="screen16x9"/>
<p:notesSz cx="6858000" cy="9144000"/>
<p:sldIdLst>{slide_ids}</p:sldIdLst>
</p:presentation>""",
        )
        rels = "\n".join(
            [f'<Relationship Id="rId{i}" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide" Target="slides/slide{i}.xml"/>' for i in range(1, len(SLIDES) + 1)]
        )
        z.writestr(
            "ppt/_rels/presentation.xml.rels",
            f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">{rels}</Relationships>""",
        )
        for i, (title, bullets) in enumerate(SLIDES, 1):
            z.writestr(f"ppt/slides/slide{i}.xml", slide_xml(title, bullets, i))


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_xlsx()
    build_pptx()
    print(f"Created assets in {OUT}")


if __name__ == "__main__":
    main()
